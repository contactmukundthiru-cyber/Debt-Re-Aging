"""
Word (DOCX) export functionality for dispute packets.
Provides high-fidelity document generation for professional editing.
"""

import os
from pathlib import Path
from typing import Dict, List, Any, Optional
import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def export_to_docx(md_content: str, output_path: str, title: str = "Dispute Document"):
    """
    Convert Markdown to a professional Word document.
    """
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Convert MD to HTML first for easier structure traversal
    html = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])
    soup = BeautifulSoup(html, 'html.parser')
    
    for element in soup.contents:
        if element.name == 'h1':
            p = doc.add_heading(element.get_text(), level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif element.name == 'h2':
            doc.add_heading(element.get_text(), level=1)
        elif element.name == 'h3':
            doc.add_heading(element.get_text(), level=2)
        elif element.name == 'p':
            p = doc.add_paragraph(element.get_text())
        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(li.get_text(), style='List Bullet')
        elif element.name == 'ol':
            for li in element.find_all('li'):
                doc.add_paragraph(li.get_text(), style='List Number')
        elif element.name == 'table':
            rows = element.find_all('tr')
            if not rows: continue
            
            # Count columns
            cols = len(rows[0].find_all(['td', 'th']))
            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = 'Table Grid'
            
            for i, row in enumerate(rows):
                cells = row.find_all(['td', 'th'])
                for j, cell in enumerate(cells):
                    table.cell(i, j).text = cell.get_text()
        elif element.name == 'blockquote':
            p = doc.add_paragraph(element.get_text())
            p.paragraph_format.left_indent = Inches(0.5)
            # Add a subtle italic style
            p.runs[0].italic = True
            
    doc.save(output_path)
    return True

def export_packet_to_docx(generated_files: Dict[str, str], output_dir: str) -> Dict[str, str]:
    """
    Export all markdown files in a packet to DOCX.
    """
    docx_files = {}
    output_path = Path(output_dir)
    
    for filename, filepath in generated_files.items():
        if not filename.endswith('.md'):
            continue
            
        with open(filepath, 'r', encoding='utf-8') as f:
            md_content = f.read()
            
        base_name = filename.replace('.md', '')
        title = base_name.replace('_', ' ').title()
        
        docx_filename = f"{base_name}.docx"
        docx_path = output_path / docx_filename
        
        try:
            if export_to_docx(md_content, str(docx_path), title):
                docx_files[docx_filename] = str(docx_path)
        except Exception as e:
            # Silently fail for now, though in a real app we'd log this
            print(f"Error exporting {filename} to DOCX: {e}")
            
    return docx_files
